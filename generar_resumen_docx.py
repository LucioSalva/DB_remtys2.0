"""
Genera Resumen_remtys_db.docx con el resumen, flujo y diagrama del proyecto.
No modifica nada del repositorio: solo escribe el .docx en la raiz del proyecto.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\lua22\Desktop\creacionSoftware\BasedeDatosRemetys2.0\Resumen_remtys_db.docx"

doc = Document()

# Margenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def add_paragraph(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_mono_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Sombreado gris claro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

# ============================================================
# PORTADA / TITULO
# ============================================================
title = doc.add_heading('Proyecto remtys_db', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Repositorio documental unico por contribuyente — Tesoreria Municipal de Ecatepec de Morelos')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()

# ============================================================
# 1. RESUMEN
# ============================================================
add_heading('1. Resumen del proyecto', level=1)

add_paragraph(
    'remtys_db es una base de datos PostgreSQL 17 que implementa un repositorio documental unico '
    'por contribuyente. La idea central: el contribuyente sube cada documento UNA SOLA VEZ; cuando '
    'inicia un nuevo tramite, el sistema solo le pide los documentos que le faltan o que ya '
    'expiraron, porque los 20 tramites municipales comparten muchos documentos comunes '
    '(identificacion, escritura, recibo predial, poder notarial, etc.).'
)

add_heading('Stack actual', level=2)
add_bullet('PostgreSQL 17 — localhost:5432, base remtys_db, usuario postgres.')
add_bullet('Sin API. Solo capa de datos: DDL, seeds, patches, vistas y triggers.')
add_bullet('Soporte: scripts Python con RapidOCR + PyMuPDF para procesar los PDFs originales.')

add_heading('Componentes del repositorio', level=2)

tabla = doc.add_table(rows=1, cols=2)
tabla.style = 'Light Grid Accent 1'
hdr = tabla.rows[0].cells
hdr[0].text = 'Carpeta / archivo'
hdr[1].text = 'Proposito'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

filas = [
    ('db/01_schema.sql',           'DDL: 11 tablas, 6 enums, 2 vistas, 1 trigger de scope.'),
    ('db/02_seed_catalogos.sql',   '32 tipo_documento + alias + grupo ACREDITA_PROPIEDAD.'),
    ('db/03_seed_tramites.sql',    '20 tramites + relaciones entre tramites.'),
    ('db/04_seed_requisitos.sql',  '309 filas de matriz tramite_requisito.'),
    ('db/05_patch_ajustes.sql',    'Ajustes idempotentes: T18 (OR), JURIDICA en T1/T6/T9/T11, anticipo_pct 20% en Convenio. Total: 342 requisitos.'),
    ('extracted_txt/*.txt',        '27 fichas estructuradas de los tramites.'),
    ('ocr_pdfs.py + ocr_raw/',     'Pipeline OCR (RapidOCR + PyMuPDF a 250 DPI).'),
    ('backups/*.sql y *.dump',     'Respaldo plano (pg_dump) y custom (pg_dump -Fc).'),
    ('remtysTesoreria/*.pdf',      '20 PDFs originales de los tramites.'),
]
for nombre, prop in filas:
    row = tabla.add_row().cells
    row[0].text = nombre
    row[1].text = prop

doc.add_paragraph()

add_heading('Modelo conceptual clave', level=2)
add_bullet('documento.scope (CONTRIBUYENTE / INMUEBLE / SOLICITUD): un trigger valida que la FK correcta este presente segun el alcance del documento.')
add_bullet('Deduplicacion: UNIQUE (contribuyente_id, sha256) en documento evita resubidas.')
add_bullet('Requisitos "alguno de": tramite_requisito apunta a tipo_documento_id XOR grupo_id (ej. T18 — Pago Predial: identificacion oficial O recibo predial).')
add_bullet('Reuso real validado: IDENTIFICACION_OFICIAL aparece en 17/20 tramites; PODER_NOTARIAL en 15/20; grupo ACREDITA_PROPIEDAD en 16/20.')
add_bullet('Vistas: v_solicitud_checklist (que falta por solicitud) y v_documentos_reutilizables (que tiene vigente el contribuyente).')

# ============================================================
# 2. FLUJO DE TRABAJO
# ============================================================
add_heading('2. Flujo de trabajo (operativo)', level=1)

pasos = [
    ('Alta del contribuyente',
     'Se crea fila en contribuyente con tipo_persona (FISICA / JURIDICA / INSTITUCION_PUBLICA).'),
    ('Subida de documento',
     'Se inserta fila en documento con sha256, vigencia_hasta, scope y la FK correspondiente '
     '(contribuyente_id, inmueble_id o solicitud_id segun scope). El trigger valida la coherencia.'),
    ('Inicio de tramite',
     'Se crea solicitud en estado BORRADOR referenciando tramite_id + contribuyente_id '
     '(+ inmueble_id cuando aplica).'),
    ('Calculo de checklist',
     'La vista v_solicitud_checklist cruza tramite_requisito con los documento vigentes del '
     'contribuyente: si existe vigente -> REUTILIZADO; si falta o expiro -> PENDIENTE.'),
    ('Carga incremental',
     'Solo se suben los pendientes. Cada nuevo documento queda en el repositorio y disponible '
     'para futuros tramites del mismo contribuyente.'),
    ('Asociacion',
     'Cada documento usado se vincula a la solicitud via solicitud_documento.'),
    ('Anticipo (T8 Convenio)',
     'tramite.anticipo_pct = 20% se valida fuera del checklist documental (es condicion de pago, no documento).'),
    ('Cambio de estado',
     'solicitud.estado avanza: BORRADOR -> EN_REVISION -> APROBADA / RECHAZADA -> CONCLUIDA.'),
]
for i, (titulo, desc) in enumerate(pasos, start=1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(f'{titulo}. ')
    r.bold = True
    p.add_run(desc)

# ============================================================
# 3. DIAGRAMA DE FLUJO
# ============================================================
add_heading('3. Diagrama de flujo', level=1)

diagrama = r"""
                       +--------------------------+
                       |   Contribuyente (alta)   |
                       +------------+-------------+
                                    |
                                    v
                       +--------------------------+
                       |   Sube documento (PDF)   |
                       |   -> calcula SHA256      |
                       +------------+-------------+
                                    |
                          +---------+----------+
                          | SHA256 ya existe?  |
                          +---------+----------+
                            si  +---+---+  no
                                v       v
                       Reusa fila   INSERT documento
                                    (scope + FK + vigencia)
                                          |
                                          v
                                 Trigger valida scope
                                          |
                                          v
                       +--------------------------+
                       |   Solicitar tramite N    |
                       |   INSERT solicitud       |
                       +------------+-------------+
                                    |
                                    v
                       +--------------------------+
                       | v_solicitud_checklist    |
                       | JOIN tramite_requisito x |
                       | documento vigente        |
                       +------------+-------------+
                                    |
                       +------------+------------+
                       v                         v
              +----------------+       +------------------+
              | Doc vigente ya |       |  Falta / vencido |
              | en repositorio |       |                  |
              +-------+--------+       +--------+---------+
                      |                         |
                      v                         v
              REUTILIZADO              Pide subir SOLO lo faltante
              (link en                          |
              solicitud_documento)              v
                      |                  Vuelve al paso "Sube documento"
                      +------------+------------+
                                   v
                       +--------------------------+
                       | Checklist completo?      |
                       +------------+-------------+
                                    | si
                                    v
                       +--------------------------+
                       | Tramite T8 (Convenio)?   |
                       | -> exige anticipo 20%    |
                       +------------+-------------+
                                    |
                                    v
                       +--------------------------+
                       | solicitud.estado:        |
                       | BORRADOR -> EN_REVISION  |
                       |  -> APROBADA/RECHAZADA   |
                       |  -> CONCLUIDA            |
                       +--------------------------+
"""
add_mono_block(diagrama)

# ============================================================
# 4. NOTAS FINALES
# ============================================================
add_heading('4. Notas finales', level=1)
add_bullet('Repositorio publicado en https://github.com/LucioSalva/DB_remtys2.0')
add_bullet('Backups validados restaurando a base temporal y comparando conteos de filas.')
add_bullet('No existe API ni capa de aplicacion: el proyecto vive integramente en PostgreSQL.')
add_bullet('Documento generado automaticamente a partir del estado actual del repositorio; no se modifico ningun archivo del proyecto.')

doc.save(OUT)
print(f'OK: {OUT}')
